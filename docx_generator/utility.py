#!/usr/bin/env python3
"""Utilities class with helper methods for creating a Word doc report.
The functions in this class are currently limited to CSF (CMMI) assessments."""

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # pylint: disable=no-name-in-module
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

OBJECTIVE_DATA_FOR_HEATMAPS = {
    "DE.AE": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "DE.CM": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "DE.DP": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "ID.AM": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "ID.BE": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "ID.GV": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "ID.RA": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "ID.RM": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "ID.SC": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "PR.AC": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "PR.AT": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "PR.DS": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "PR.IP": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "PR.MA": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "PR.PT": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "RC.CO": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "RC.IM": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "RC.RP": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "RS.AN": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "RS.CO": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "RS.IM": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "RS.MI": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
    "RS.RP": {
        "average_practice_level": 0,
        "objective_text": "",
        "practices": [],
    },
}


def _set_table_cell_properties(table_cell, color_hex, **kwargs):
    if color_hex is None:
        return

    # pylint: disable=protected-access
    table_cell = table_cell._tc
    table_cell_properties = table_cell.get_or_add_tcPr()
    cell_shading = OxmlElement("w:shd")
    cell_shading.set(qn("w:fill"), color_hex)
    table_cell_properties.append(cell_shading)

    # check for cell borders tag existence, if none found, then create one
    table_cell_borders = table_cell_properties.first_child_found_in("w:tcBorders")
    if table_cell_borders is None:
        table_cell_borders = OxmlElement('w:tcBorders')
        table_cell_properties.append(table_cell_borders)

    # list over all available border tags
    for edge in kwargs:
        edge_data = kwargs.get(edge)
        tag = 'w:{}'.format(edge)

        # check for given tag existence, if none found, then create one
        element = table_cell_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            table_cell_borders.append(element)

            for key in edge_data:
                element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def _set_table_cell_width(table_cell, paragraph):
    paragraph_elements = paragraph.text.split("#|#")

    number_of_practices = int(paragraph_elements[0])
    total_practices_in_row = int(paragraph_elements[1])
    paragraph.text = paragraph_elements[0]
    for run in paragraph.runs:
        run.font.bold = False

    if number_of_practices == 0:
        paragraph.text = ""

    # pylint: disable=protected-access
    table_cell = table_cell._tc
    table_cell_properties = table_cell.get_or_add_tcPr()
    if number_of_practices is not None:
        for prop in table_cell_properties:
            if prop.tag == qn('w:tcW'):
                table_cell_properties.remove(prop)

    width_in_inches = number_of_practices if number_of_practices == 0 else (
            number_of_practices / total_practices_in_row * 8)

    table_cell.width = Inches(width_in_inches)


# hardcoded levels from the cmmicsf model file
LEVEL_NAMES = {
    "-1": "Unset",
    0: "Incomplete",
    1: "Initial",
    2: "Managed",
    3: "Defined",
    4: "Quantitatively Managed",
    5: "Optimizing",
}

TABLE_CELL_COLORS = {
    "Unset": None,
    "Incomplete": "FF001C",
    "Initial": "FFBE37",
    "Managed": "FFFD46",
    "Defined": "91CF5D",
    "Quantitatively Managed": "6C9E25",
    "Optimizing": "1D451B",
}

DOMAIN_DATA_FOR_HEATMAPS = {
    "ID": {"domain_text": "", "average_practice_level": 0},
    "PR": {"domain_text": "", "average_practice_level": 0},
    "DE": {"domain_text": "", "average_practice_level": 0},
    "RS": {"domain_text": "", "average_practice_level": 0},
    "RC": {"domain_text": "", "average_practice_level": 0},
}

GROUPED_BY_LEVEL = {"-1": [], "0": [], "1": [], "2": [], "3": [], "4": [], "5": []}


def add_notes_to_table_cells(results_doc):
    """Take the doc created via MailMerge, search for Notes
    mergefields and reposition/reformat them, then save the changes"""
    document = Document(results_doc)
    for table in document.tables:
        for column in table.columns:
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    _style_notes_in_paragraph(paragraph)
    document.save(results_doc)


def _style_notes_in_paragraph(paragraph):
    for note in ALL_NOTES:
        if note['id'] in paragraph.text:
            note_regex = re.compile(note['id'])
            new_line_regex = re.compile('\n')
            if note_regex.search(paragraph.text):
                paragraph.text = note_regex.sub('', paragraph.text)
            if new_line_regex.search(note["text"]):
                text_split = note["text"].split("\n")
                for note_part in text_split:
                    para = paragraph.insert_paragraph_before()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph_format = para.paragraph_format
                    paragraph_format.space_before = Pt(6)
                    new_note = para.add_run(note_part)
                    new_note.font.size = Pt(9)
            else:
                para = paragraph.insert_paragraph_before()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph_format = para.paragraph_format
                paragraph_format.space_before = Pt(6)
                new_note = para.add_run(note['text'])
                new_note.font.size = Pt(9)


def add_critical_practice_to_table_cells(results_doc):
    """Take the doc created via MailMerge, search for Critical Control
    mergefields and reposition/reformat them, then save the changes"""
    document = Document(results_doc)
    for table in document.tables:
        for column in table.columns:
            for cell in column.cells:
                for paragraph in cell.paragraphs:
                    note_regex = re.compile("Critical_Control")
                    if note_regex.search(paragraph.text):
                        paragraph.text = note_regex.sub('', paragraph.text)
                        para = cell.paragraphs[0].insert_paragraph_before()
                        paragraph_format = para.paragraph_format
                        paragraph_format.space_before = Pt(6)
                        new_critical_practice_line = para.add_run("Critical Control")
                        new_critical_practice_line.font.size = Pt(9)
                        new_critical_practice_line.bold = True
                        new_critical_practice_line.font.color.rgb = RGBColor.from_string('FF0000')
    document.save(results_doc)


def style_table_cells(results_doc):
    """Take the doc created via MailMerge and color every cell of every table
    according to what text the cell contains, then save the changes"""
    document = Document(results_doc)
    for table in document.tables:
        for column in table.columns:
            for cell in column.cells:
                _style_cell_according_to_text_inside(cell)
    document.save(results_doc)


def _style_cell_according_to_text_inside(cell):
    paragraphs_to_check = [
        paragraph
        for paragraph in cell.paragraphs
        if len(cell.paragraphs) <= 2
    ]

    cell_is_blank = True
    for paragraph in paragraphs_to_check:
        if paragraph.text != "":
            cell_is_blank = False

    for paragraph in paragraphs_to_check:
        for level, group in GROUPED_BY_LEVEL.items():
            if paragraph.text in group:
                level_text = LEVEL_NAMES.get(int(level))
                background_color = TABLE_CELL_COLORS.get(level_text)
                border_color = "auto"
            elif cell_is_blank or "0#|#" in paragraph.text:
                border_color = "#FFFFFF"
                background_color = "#FFFFFF"
            else:
                border_color = "auto"
                background_color = TABLE_CELL_COLORS.get(paragraph.text)

            _set_table_cell_properties(cell, background_color,
                                       start={"val": "single", "color": border_color},
                                       top={"val": "single", "color": border_color},
                                       end={"val": "single", "color": border_color},
                                       bottom={"val": "single", "color": border_color})

            if "#|#" in paragraph.text:
                _set_table_cell_width(cell, paragraph)


ID_ROW_DATA = []
PR_ROW_DATA = []
DE_ROW_DATA = []
RS_ROW_DATA = []
RC_ROW_DATA = []

ROW_SWITCHER = {
    "ID": ID_ROW_DATA,
    "PR": PR_ROW_DATA,
    "DE": DE_ROW_DATA,
    "RS": RS_ROW_DATA,
    "RC": RC_ROW_DATA,
}

ID_HEATMAP_ROW_DATA = []
PR_HEATMAP_ROW_DATA = []
DE_HEATMAP_ROW_DATA = []
RS_HEATMAP_ROW_DATA = []
RC_HEATMAP_ROW_DATA = []

OBJECTIVE_HEATMAP_ROW_SWITCHER = {
    "ID": ID_HEATMAP_ROW_DATA,
    "PR": PR_HEATMAP_ROW_DATA,
    "DE": DE_HEATMAP_ROW_DATA,
    "RS": RS_HEATMAP_ROW_DATA,
    "RC": RC_HEATMAP_ROW_DATA,
}

ALL_NOTES = []


def merge_objective_heatmap_rows(template):
    """Take the MailMerge template object and merge the heatmap data for each objective/category"""
    template.merge_rows(
        "ID.AM_objective",
        [item for item in ID_HEATMAP_ROW_DATA if "ID.AM_objective" in item],
    )
    template.merge_rows(
        "ID.BE_objective",
        [item for item in ID_HEATMAP_ROW_DATA if "ID.BE_objective" in item],
    )
    template.merge_rows(
        "ID.GV_objective",
        [item for item in ID_HEATMAP_ROW_DATA if "ID.GV_objective" in item],
    )
    template.merge_rows(
        "ID.RA_objective",
        [item for item in ID_HEATMAP_ROW_DATA if "ID.RA_objective" in item],
    )
    template.merge_rows(
        "ID.RM_objective",
        [item for item in ID_HEATMAP_ROW_DATA if "ID.RM_objective" in item],
    )
    template.merge_rows(
        "ID.SC_objective",
        [item for item in ID_HEATMAP_ROW_DATA if "ID.SC_objective" in item],
    )
    template.merge_rows(
        "PR.AC_objective",
        [item for item in PR_HEATMAP_ROW_DATA if "PR.AC_objective" in item],
    )
    template.merge_rows(
        "PR.AT_objective",
        [item for item in PR_HEATMAP_ROW_DATA if "PR.AT_objective" in item],
    )
    template.merge_rows(
        "PR.DS_objective",
        [item for item in PR_HEATMAP_ROW_DATA if "PR.DS_objective" in item],
    )
    template.merge_rows(
        "PR.IP_objective",
        [item for item in PR_HEATMAP_ROW_DATA if "PR.IP_objective" in item],
    )
    template.merge_rows(
        "PR.MA_objective",
        [item for item in PR_HEATMAP_ROW_DATA if "PR.MA_objective" in item],
    )
    template.merge_rows(
        "PR.PT_objective",
        [item for item in PR_HEATMAP_ROW_DATA if "PR.PT_objective" in item],
    )
    template.merge_rows(
        "DE.AE_objective",
        [item for item in DE_HEATMAP_ROW_DATA if "DE.AE_objective" in item],
    )
    template.merge_rows(
        "DE.CM_objective",
        [item for item in DE_HEATMAP_ROW_DATA if "DE.CM_objective" in item],
    )
    template.merge_rows(
        "DE.DP_objective",
        [item for item in DE_HEATMAP_ROW_DATA if "DE.DP_objective" in item],
    )
    template.merge_rows(
        "RS.AN_objective",
        [item for item in RS_HEATMAP_ROW_DATA if "RS.AN_objective" in item],
    )
    template.merge_rows(
        "RS.CO_objective",
        [item for item in RS_HEATMAP_ROW_DATA if "RS.CO_objective" in item],
    )
    template.merge_rows(
        "RS.IM_objective",
        [item for item in RS_HEATMAP_ROW_DATA if "RS.IM_objective" in item],
    )
    template.merge_rows(
        "RS.MI_objective",
        [item for item in RS_HEATMAP_ROW_DATA if "RS.MI_objective" in item],
    )
    template.merge_rows(
        "RS.RP_objective",
        [item for item in RS_HEATMAP_ROW_DATA if "RS.RP_objective" in item],
    )
    template.merge_rows(
        "RC.CO_objective",
        [item for item in RC_HEATMAP_ROW_DATA if "RC.CO_objective" in item],
    )
    template.merge_rows(
        "RC.IM_objective",
        [item for item in RC_HEATMAP_ROW_DATA if "RC.IM_objective" in item],
    )
    template.merge_rows(
        "RC.RP_objective",
        [item for item in RC_HEATMAP_ROW_DATA if "RC.RP_objective" in item],
    )
